from pathlib import Path
from collections import Counter
from zipfile import ZipFile
from docx import Document
from pypdf import PdfReader
import hashlib
import json

ROOT = Path(r"D:\SistemaHidrológico")
DOCX = ROOT / "docs" / "Documento_referencia_PAE.docx"
PDF = ROOT / "docs" / "_qa" / "reference" / "reference.pdf"
OUT = ROOT / "docs" / "_qa" / "reference"

doc = Document(DOCX)
paragraphs = []
for i, p in enumerate(doc.paragraphs):
    text = p.text.strip()
    if text:
        paragraphs.append({"index": i, "style": p.style.name if p.style else "", "text": text})

tables = []
for ti, table in enumerate(doc.tables):
    sample = []
    for row in table.rows[:5]:
        sample.append([cell.text.strip().replace("\n", " | ")[:250] for cell in row.cells])
    tables.append({"index": ti, "rows": len(table.rows), "cols": len(table.columns), "sample": sample})

reader = PdfReader(str(PDF))
page_texts = []
for i, page in enumerate(reader.pages):
    try:
        text = page.extract_text() or ""
    except Exception:
        text = ""
    page_texts.append(f"\n\n===== PAGE {i + 1} =====\n{text}")

with ZipFile(DOCX) as zf:
    parts = [{"path": x.filename, "size": x.file_size, "sha256": hashlib.sha256(zf.read(x.filename)).hexdigest()} for x in zf.infolist()]

summary = {
    "sha256": hashlib.sha256(DOCX.read_bytes()).hexdigest(),
    "pages": len(reader.pages),
    "sections": len(doc.sections),
    "paragraph_count": len(doc.paragraphs),
    "nonempty_paragraph_count": len(paragraphs),
    "table_count": len(doc.tables),
    "inline_shape_count": len(doc.inline_shapes),
    "style_usage": Counter(p["style"] for p in paragraphs).most_common(),
    "paragraphs": paragraphs,
    "tables": tables,
    "package_parts": parts,
}
(OUT / "inspection.json").write_text(json.dumps(summary, ensure_ascii=False, indent=2), encoding="utf-8")
(OUT / "reference_text.txt").write_text("".join(page_texts), encoding="utf-8")
print(json.dumps({k: summary[k] for k in ("sha256", "pages", "sections", "paragraph_count", "table_count", "inline_shape_count", "style_usage")}, ensure_ascii=False, indent=2))
