from pathlib import Path
from datetime import date
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_LINE_SPACING
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.section import WD_SECTION
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.enum.style import WD_STYLE_TYPE

ROOT = Path(__file__).resolve().parents[1]
OUT = ROOT / "docs" / "Elicitacion_de_Requisitos_Dashboard_Hidrologico_de_Prevencion.docx"

NAVY = "08232D"
DEEP = "0B3340"
TEAL = "008F86"
CYAN = "2DCFC2"
PALE = "E8F5F4"
LIGHT = "F2F6F7"
MID = "607D86"
INK = "17333B"
WHITE = "FFFFFF"
GOLD = "C58A20"
RED = "B13A45"

doc = Document()
sec = doc.sections[0]
sec.page_width, sec.page_height = Inches(8.5), Inches(11)
sec.top_margin = sec.bottom_margin = sec.left_margin = sec.right_margin = Inches(1)
sec.header_distance = sec.footer_distance = Inches(.48)
doc.core_properties.title = "Dashboard Hidrológico de Prevención - Elicitación de Requisitos"
doc.core_properties.subject = "Especificación integral del dashboard preventivo multi-río con datos IoT de ThingSpeak"
doc.core_properties.author = "Equipo HydroWatch Perú"
doc.core_properties.keywords = "ThingSpeak, IoT, hidrología, dashboard, requisitos"

styles = doc.styles

def set_font(style, name="Aptos", size=11, color=INK, bold=None):
    style.font.name = name
    style._element.rPr.rFonts.set(qn("w:ascii"), name)
    style._element.rPr.rFonts.set(qn("w:hAnsi"), name)
    style.font.size = Pt(size)
    style.font.color.rgb = RGBColor.from_string(color)
    if bold is not None: style.font.bold = bold

normal = styles["Normal"]
set_font(normal, size=10.5)
normal.paragraph_format.space_after = Pt(6)
normal.paragraph_format.line_spacing = 1.16
normal.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

for name, size, color, before, after in [
    ("Title", 28, NAVY, 0, 6), ("Subtitle", 13, MID, 0, 12),
    ("Heading 1", 17, TEAL, 16, 8), ("Heading 2", 13, DEEP, 12, 6), ("Heading 3", 11, TEAL, 9, 4)
]:
    st = styles[name]; set_font(st, size=size, color=color, bold=name != "Subtitle")
    st.paragraph_format.space_before = Pt(before); st.paragraph_format.space_after = Pt(after)
    st.paragraph_format.keep_with_next = True
    if name == "Heading 1": st.paragraph_format.page_break_before = True

for list_name in ("List Bullet", "List Number"):
    st = styles[list_name]; set_font(st, size=10.5)
    st.paragraph_format.left_indent = Inches(.5); st.paragraph_format.first_line_indent = Inches(-.25)
    st.paragraph_format.space_after = Pt(4); st.paragraph_format.line_spacing = 1.15

if "Requirement ID" not in styles:
    req_style = styles.add_style("Requirement ID", WD_STYLE_TYPE.PARAGRAPH)
    set_font(req_style, size=9, color=TEAL, bold=True)
    req_style.paragraph_format.space_after = Pt(2)

def shade(cell, fill):
    tcPr = cell._tc.get_or_add_tcPr(); shd = tcPr.find(qn("w:shd"))
    if shd is None: shd = OxmlElement("w:shd"); tcPr.append(shd)
    shd.set(qn("w:fill"), fill)

def set_cell_margins(cell, top=90, start=110, bottom=90, end=110):
    tc = cell._tc; tcPr = tc.get_or_add_tcPr(); tcMar = tcPr.first_child_found_in("w:tcMar")
    if tcMar is None: tcMar = OxmlElement("w:tcMar"); tcPr.append(tcMar)
    for edge, value in (("top",top),("start",start),("bottom",bottom),("end",end)):
        node = tcMar.find(qn(f"w:{edge}"))
        if node is None: node = OxmlElement(f"w:{edge}"); tcMar.append(node)
        node.set(qn("w:w"), str(value)); node.set(qn("w:type"), "dxa")

def set_repeat_table_header(row):
    trPr = row._tr.get_or_add_trPr(); tag = OxmlElement("w:tblHeader"); tag.set(qn("w:val"), "true"); trPr.append(tag)

def prevent_row_split(row):
    trPr = row._tr.get_or_add_trPr()
    if trPr.find(qn("w:cantSplit")) is None:
        trPr.append(OxmlElement("w:cantSplit"))

def set_table_geometry(table, widths):
    total = sum(widths); table.autofit = False; table.alignment = WD_TABLE_ALIGNMENT.LEFT
    tblPr = table._tbl.tblPr
    for tag_name, attrs in (("tblW", {"w": total, "type": "dxa"}), ("tblInd", {"w": 110, "type": "dxa"}), ("tblLayout", {"type": "fixed"})):
        old = tblPr.find(qn(f"w:{tag_name}"))
        if old is not None: tblPr.remove(old)
        node = OxmlElement(f"w:{tag_name}")
        for k,v in attrs.items(): node.set(qn(f"w:{k}"), str(v))
        tblPr.append(node)
    grid = table._tbl.tblGrid
    for child in list(grid): grid.remove(child)
    for width in widths:
        col = OxmlElement("w:gridCol"); col.set(qn("w:w"), str(width)); grid.append(col)
    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            cell.width = Inches(widths[idx]/1440)
            tcPr = cell._tc.get_or_add_tcPr(); tcW = tcPr.find(qn("w:tcW"))
            tcW.set(qn("w:w"), str(widths[idx])); tcW.set(qn("w:type"), "dxa")
            set_cell_margins(cell); cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER

def table(headers, rows, widths, font_size=8.5):
    t = doc.add_table(rows=1, cols=len(headers)); t.style = "Table Grid"
    for i, h in enumerate(headers):
        cell=t.rows[0].cells[i]; cell.text=str(h); shade(cell, DEEP)
        for p in cell.paragraphs:
            p.alignment=WD_ALIGN_PARAGRAPH.CENTER; p.paragraph_format.space_after=Pt(0)
            for r in p.runs: r.font.bold=True; r.font.color.rgb=RGBColor(255,255,255); r.font.size=Pt(font_size)
    set_repeat_table_header(t.rows[0])
    prevent_row_split(t.rows[0])
    for ri, values in enumerate(rows):
        cells=t.add_row().cells
        for i, value in enumerate(values):
            cells[i].text=str(value)
            if ri % 2: shade(cells[i], LIGHT)
            for p in cells[i].paragraphs:
                p.paragraph_format.space_after=Pt(1); p.paragraph_format.line_spacing=1.05
                for r in p.runs: set_run(r, size=font_size, color=INK)
        prevent_row_split(t.rows[-1])
    set_table_geometry(t, widths)
    doc.add_paragraph().paragraph_format.space_after=Pt(0)
    return t

def set_run(run, size=10.5, color=INK, bold=False, italic=False, name="Aptos"):
    run.font.name=name; run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"),name); run._element.rPr.rFonts.set(qn("w:hAnsi"),name)
    run.font.size=Pt(size); run.font.color.rgb=RGBColor.from_string(color); run.bold=bold; run.italic=italic

def add_p(text="", bold_lead=None, align=None, after=6):
    p=doc.add_paragraph(); p.paragraph_format.space_after=Pt(after)
    if align is not None: p.alignment=align
    if bold_lead and text.startswith(bold_lead):
        a=p.add_run(bold_lead); set_run(a,bold=True); b=p.add_run(text[len(bold_lead):]); set_run(b)
    else: r=p.add_run(text); set_run(r)
    return p

def bullet(text):
    p=doc.add_paragraph(style="List Bullet"); p.add_run(text); return p

def heading(text, level=1): return doc.add_heading(text, level=level)

def callout(label, text, fill=PALE, accent=TEAL):
    t=doc.add_table(rows=1, cols=1); t.style="Table Grid"; c=t.cell(0,0); shade(c,fill); set_cell_margins(c,150,180,150,180)
    p=c.paragraphs[0]; p.paragraph_format.space_after=Pt(2); a=p.add_run(label.upper()+"  "); set_run(a,size=9,color=accent,bold=True); b=p.add_run(text); set_run(b,size=10,color=INK)
    set_repeat_table_header(t.rows[0]); prevent_row_split(t.rows[0]); set_table_geometry(t,[9360]); doc.add_paragraph().paragraph_format.space_after=Pt(0)

def add_page_field(paragraph):
    run=paragraph.add_run(); fldChar=OxmlElement("w:fldChar"); fldChar.set(qn("w:fldCharType"),"begin"); instr=OxmlElement("w:instrText"); instr.set(qn("xml:space"),"preserve"); instr.text=" PAGE "; sep=OxmlElement("w:fldChar"); sep.set(qn("w:fldCharType"),"separate"); text=OxmlElement("w:t"); text.text="1"; end=OxmlElement("w:fldChar"); end.set(qn("w:fldCharType"),"end")
    for node in (fldChar,instr,sep,text,end): run._r.append(node)

def setup_header_footer(section, first=False):
    header=section.header; p=header.paragraphs[0]; p.clear(); p.text="DASHBOARD HIDROLÓGICO DE PREVENCIÓN  |  ELICITACIÓN DE REQUISITOS"; p.alignment=WD_ALIGN_PARAGRAPH.RIGHT
    for r in p.runs: set_run(r,size=8,color=MID,bold=True)
    footer=section.footer; p=footer.paragraphs[0]; p.clear(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER
    r=p.add_run("Documento de especificación  •  v1.1  •  Página "); set_run(r,size=8,color=MID); add_page_field(p)

setup_header_footer(sec)
settings = doc.settings._element
update = OxmlElement("w:updateFields"); update.set(qn("w:val"),"true"); settings.append(update)

# Portada
p=doc.add_paragraph(); p.paragraph_format.space_before=Pt(42); r=p.add_run("HYDROWATCH PERÚ"); set_run(r,size=11,color=TEAL,bold=True); p.alignment=WD_ALIGN_PARAGRAPH.CENTER
p=doc.add_paragraph(); p.paragraph_format.space_before=Pt(52); p.alignment=WD_ALIGN_PARAGRAPH.CENTER; r=p.add_run("DASHBOARD HIDROLÓGICO\nDE PREVENCIÓN"); set_run(r,size=29,color=NAVY,bold=True)
p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER; r=p.add_run("Documento de Elicitación y Especificación de Requisitos"); set_run(r,size=14,color=MID)
p=doc.add_paragraph(); p.paragraph_format.space_before=Pt(30); p.alignment=WD_ALIGN_PARAGRAPH.CENTER; r=p.add_run("Red multi-río  •  Estación inicial: Río Huallaga"); set_run(r,size=11,color=TEAL,bold=True)
callout("Propósito", "Establecer una base común, verificable y trazable para construir y validar el dashboard que transforma la telemetría hidrológica en información operativa clara.")
p=doc.add_paragraph(); p.paragraph_format.space_before=Pt(72); p.alignment=WD_ALIGN_PARAGRAPH.CENTER; r=p.add_run("Versión 1.1  |  22 de julio de 2026  |  Estado: Actualizado"); set_run(r,size=9,color=MID)
doc.add_page_break()

heading("CONTROL DEL DOCUMENTO",1)
table(["Versión","Fecha","Autor / responsable","Descripción","Estado"], [["1.1","22/07/2026","Equipo HydroWatch Perú","Especificación integral del Dashboard Hidrológico de Prevención, estación inicial Río Huallaga y cobertura multi-río.","Actualizado"]],[850,1150,1850,4000,1510],8.5)
heading("Aprobaciones requeridas",2)
table(["Rol","Responsabilidad de validación","Nombre / firma","Fecha"], [["Product Owner","Alcance y prioridad","Pendiente","—"],["Responsable IoT","Unidades, campos, códigos de estado y frecuencia","Pendiente","—"],["Usuario operativo","Usabilidad y lectura de alertas","Pendiente","—"],["Equipo técnico","Viabilidad, seguridad y pruebas","Pendiente","—"]],[1500,3900,2500,1460],8.5)
callout("Nota de validación", "La correspondencia de field7 se interpreta provisionalmente como 0 Normal, 1 Preventivo, 2 Alerta y 3 Crítico. El responsable del firmware o del modelo debe confirmarla antes del uso operativo.", "FFF5DF", GOLD)

heading("CONTENIDO",1)
for item in ["1. Introducción y propósito","2. Alcance del proyecto","3. Contexto del producto y fuentes de datos","4. Stakeholders y mapa de poder/interés","5. Plan, técnicas y sprints de elicitación","6. Hallazgos, supuestos y restricciones","7. Requisitos funcionales","8. Requisitos no funcionales","9. Historias de usuario y criterios de aceptación","10. Casos de uso","11. Product Backlog y priorización","12. Matriz de trazabilidad","13. Plan de validación y aceptación","14. Riesgos","15. Glosario y anexos","16. Despliegue, dominio y operación","17. Diseño detallado de la interfaz","18. Modelo multi-río y prevención","19. Implementación, DevOps y mantenimiento"]:
    bullet(item)

heading("1. INTRODUCCIÓN Y PROPÓSITO",1)
add_p("Este documento recopila, estructura y especifica los requisitos del Dashboard Hidrológico de Prevención, una solución web multi-río destinada a convertir telemetría hidrológica en información clara, trazable y útil para la observación preventiva. El canal 3420787 se identifica como la estación inicial del Río Huallaga; los demás ríos del catálogo nacional permanecen sin telemetría hasta asociarles una estación propia.")
add_p("La elicitación se basa en tres fuentes: la solicitud del patrocinador, la inspección del canal 3420787 y el patrón documental proporcionado como referencia. No se presentan entrevistas o encuestas como ejecutadas; sus resultados deben incorporarse en versiones posteriores cuando se realicen.")
heading("1.1 Objetivos",2)
for text in ["Representar cada variable con una escala que evite perder detalle por valores atípicos o unidades incompatibles.","Proporcionar una lectura inmediata del estado actual y conservar la posibilidad de analizar tendencias.","Consultar ThingSpeak sin incorporar credenciales sensibles en el código fuente.","Definir criterios de aceptación medibles y una trazabilidad completa entre necesidades, requisitos y pruebas.","Ofrecer una experiencia atractiva, responsive y comprensible para usuarios técnicos y no técnicos."]: bullet(text)
heading("1.2 Convenciones",2)
table(["Prefijo","Uso","Ejemplo"],[["NEC","Necesidad del stakeholder","NEC-01"],["RF","Requisito funcional","RF-08"],["RNF","Requisito no funcional","RNF-04"],["HU","Historia de usuario","HU-05"],["CU","Caso de uso","CU-02"],["CP","Criterio de prueba","CP-07"]],[1200,4300,3860],9)

heading("2. ALCANCE DEL PROYECTO",1)
heading("2.1 Alcance incluido",2)
for text in ["Dashboard web de solo lectura con selección de río y canal ThingSpeak por estación.","Catálogo nacional de ríos basado en la cartografía oficial IGN 1:500 000.","Identificación del canal 3420787 como telemetría del Río Huallaga.","Indicadores de nivel, lluvia, temperatura, humedad, velocidad, predicción y estado.","Gráficos históricos con ventanas de 50, 100, 250 y 500 lecturas.","Actualización automática y manual.","Tabla de lecturas recientes y exportación CSV con río, localidad y región hidrográfica.","Configuración de Read API Key en memoria de sesión para canales privados.","Diseño responsive para escritorio, tableta y móvil.","Mensajes de conectividad, estación pendiente, ausencia de datos y errores de la API."]: bullet(text)
heading("2.2 Fuera de alcance",2)
for text in ["Escritura de datos hacia ThingSpeak o control remoto del sensor.","Modificación del firmware del dispositivo IoT.","Generación o rotación automática de API Keys.","Notificaciones por SMS, correo o mensajería en la primera versión.","Autenticación multiusuario, administración de roles o persistencia en una base de datos propia.","Certificación de predicciones o decisiones automáticas de emergencia."]: bullet(text)
callout("Límite operativo", "El dashboard apoya la interpretación, pero no reemplaza protocolos oficiales de emergencia ni la inspección física de la estación.")

heading("3. CONTEXTO DEL PRODUCTO Y FUENTES DE DATOS",1)
heading("3.1 Canal observado",2)
table(["Atributo","Valor observado"],[["ID","3420787"],["Nombre","Proyecto hidrológico"],["Descripción","Proyecto de sensor hidrológico y predicciones de agua"],["Creación","02/07/2026 23:56 UTC"],["Última actualización de metadatos observada","08/07/2026 01:47 UTC"],["Último entry_id observado durante la elicitación","273"],["Modo de lectura observado","Público; preparado para Read API Key opcional"]],[2700,6660],9)
heading("3.2 Diccionario de campos ThingSpeak",2)
table(["Campo","Nombre","Tipo esperado","Unidad / dominio","Uso en el dashboard"],[["field1","Nivel","Decimal","m (por validar)","KPI y tendencia observada"],["field2","Lluvia","Decimal / bandera","mm o 0/1 (por validar)","Evento e indicador"],["field3","Temp","Decimal","°C","KPI y serie ambiental"],["field4","Hum","Decimal","% HR","KPI y serie ambiental"],["field5","Velocidad","Decimal","por validar","Dinámica del nivel"],["field6","Nivel Predicción","Decimal","m o salida de modelo (por validar)","Serie independiente"],["field7","Estado Actual","Entero","0–3 (supuesto)","Semáforo operativo"]],[950,1750,1450,1900,3310],8.2)
heading("3.3 Arquitectura lógica",2)
add_p("El navegador realiza una petición HTTPS de solo lectura a `api.thingspeak.com/channels/{channel}/feeds.json`. La respuesta JSON se normaliza, se convierte a números y fechas, se resume en indicadores y se representa con gráficos. La clave, cuando sea necesaria, se mantiene únicamente en sessionStorage y se elimina al cerrar la pestaña.")
table(["Componente","Responsabilidad","Entrada","Salida"],[["ThingSpeak","Almacenar y servir telemetría","Datos del dispositivo","JSON de canal y feeds"],["Capa de acceso","Construir URL, solicitar y validar respuesta","Canal, ventana, clave opcional","Datos o error controlado"],["Normalización","Convertir cadenas a tipos útiles","Feeds JSON","Lecturas tipadas"],["Presentación","Calcular KPI, gráficos y tabla","Lecturas tipadas","Dashboard interactivo"],["Exportación","Generar archivo local","Lecturas visibles","CSV"]],[1700,3000,2100,2560],8.5)

heading("4. STAKEHOLDERS Y MAPA DE PODER/INTERÉS",1)
table(["ID","Stakeholder","Tipo","Interés principal","Influencia","Estrategia"],[["ST-01","Product Owner / patrocinador","Primario","Entrega y valor del dashboard","Alta","Gestionar de cerca"],["ST-02","Responsable IoT / firmware","Primario","Consistencia de campos y estados","Alta","Validación técnica"],["ST-03","Operador o analista hidrológico","Primario","Interpretación rápida y confiable","Alta","Pruebas de uso"],["ST-04","Equipo de desarrollo","Interno","Requisitos claros y mantenibles","Alta","Colaboración continua"],["ST-05","Responsable de seguridad","Secundario","Protección de claves y exposición","Media","Revisión puntual"],["ST-06","Autoridad / comunidad usuaria","Secundario","Información comprensible y oportuna","Media","Mantener informada"],["ST-07","MathWorks / ThingSpeak","Externo","Disponibilidad y límites del servicio","Baja directa","Monitorear dependencia"]],[700,1800,1050,2700,950,2160],7.8)
heading("4.1 Matriz poder/interés",2)
table(["Cuadrante","Stakeholders","Tratamiento"],[["Alto poder / alto interés","ST-01, ST-02, ST-03, ST-04","Participación en decisiones y aceptación"],["Alto poder / bajo interés","Responsable institucional futuro","Mantener satisfecho con reportes breves"],["Bajo poder / alto interés","ST-06","Demostraciones y validación de comprensión"],["Bajo poder / bajo interés","ST-07","Monitorear cambios de API y disponibilidad"]],[2100,3400,3860],8.7)

heading("5. PLAN Y TÉCNICAS DE ELICITACIÓN",1)
table(["Técnica","Objetivo","Participantes","Producto esperado"],[["Entrevista semiestructurada","Confirmar unidades, umbrales y decisiones","PO, IoT, operador","Glosario y reglas validadas"],["Observación contextual","Comprender cómo se interpreta la telemetría","Operador","Flujo y problemas de uso"],["Taller de prototipo","Validar jerarquía, colores y gráficos","PO, operador, desarrollo","Cambios priorizados"],["Análisis de datos","Detectar rangos, vacíos y atípicos","IoT, desarrollo","Reglas de visualización"],["Benchmarking","Comparar patrones de dashboards ambientales","UX, desarrollo","Buenas prácticas aplicables"],["Pruebas de tarea","Medir comprensión y tiempo de lectura","Usuarios representativos","Evidencia de usabilidad"],["Revisión documental","Alinear requisitos con la estructura acordada","Equipo","Especificación trazable"]],[1700,2600,2300,2760],8.2)
heading("5.1 Guion mínimo para entrevista",2)
for q in ["¿Qué unidad y rango válido corresponde a cada campo?","¿Qué significa cada código de Estado Actual y qué acción requiere?","¿Con qué frecuencia se espera una nueva lectura?","¿Cuándo una variación de nivel se considera anómala?","¿La lluvia es una cantidad acumulada o un indicador binario?","¿Quién puede acceder a la clave de lectura si el canal se privatiza?","¿Qué período es más útil para la operación cotidiana?","¿Qué información debe exportarse y con qué propósito?","¿Qué dispositivo y resolución usa normalmente el operador?","¿Qué evidencia permitirá aceptar el dashboard?"]: bullet(q)
heading("5.2 Product Backlog de elicitación",2)
table(["ID","Actividad","Prioridad 1–10","Esfuerzo 1–10","Estado"],[["EL-01","Confirmar unidades de los siete campos","10","2","Pendiente"],["EL-02","Confirmar escala 0–3 de field7","10","2","Pendiente"],["EL-03","Validar prototipo con operador","9","5","Pendiente"],["EL-04","Definir umbrales hidrológicos","10","5","Pendiente"],["EL-05","Medir frecuencia y latencia de actualización","8","3","En análisis"],["EL-06","Validar privacidad de canal/API Key","9","3","En análisis"],["EL-07","Ejecutar prueba responsive","8","3","Preparado"],["EL-08","Aprobar criterios de aceptación","10","4","Pendiente"]],[950,4300,1250,1250,1610],8.2)
heading("5.3 Desarrollo propuesto de sprints de elicitación",2)
add_p("El documento de referencia organiza la elicitación mediante sprints. Para este proyecto se conserva ese patrón como plan verificable; las actividades no se consideran ejecutadas hasta adjuntar evidencia y aprobación del responsable indicado.")
table(["Sprint","Objetivo","Actividades principales","Entregable","Responsable de validar"],[["1","Preparar la investigación","Identificar actores, preparar entrevistas y revisar ThingSpeak","Guiones y mapa de stakeholders","Product Owner"],["2","Comprender la operación","Entrevistar a responsable IoT y operador; observar interpretación de lecturas","Hallazgos y diccionario preliminar","Responsable IoT"],["3","Validar necesidades","Aplicar encuesta breve y priorizar problemas de comprensión","Matriz de necesidades","Usuario operativo"],["4","Explorar soluciones","Brainstorming y comparación de dashboards ambientales","Alternativas y decisiones UX","Product Owner"],["5","Validar prototipo","Recorrer resumen, gráficos, directorio y tabla con tareas reales","Registro de observaciones","Operador"],["6","Consolidar requisitos","Clasificar RF/RNF, eliminar duplicados y definir criterios medibles","Backlog depurado","Equipo técnico"],["7","Cerrar aceptación","Trazar necesidades, requisitos, historias y pruebas; resolver pendientes críticos","Línea base para aprobación","Comité del proyecto"]],[900,1700,2850,2050,1860],7.7)

heading("6. HALLAZGOS, SUPUESTOS Y RESTRICCIONES",1)
heading("6.1 Hallazgos observados",2)
for text in ["El canal expone siete campos y al menos 273 entradas al momento de la inspección.","Las variables tienen magnitudes diferentes; la predicción y la velocidad pueden presentar picos que aplanan el nivel real en un gráfico compartido.","Existen códigos de estado 0, 1 y 3 en la muestra analizada; su semántica debe confirmarse.","La lectura pública funciona sin clave, aunque el flujo debe soportar un canal privado.","El muestreo observado incluye intervalos cercanos a 20 segundos y períodos sin lecturas.","Los valores atípicos deben mostrarse sin ocultarse ni recortarse silenciosamente."]: bullet(text)
heading("6.2 Supuestos pendientes de validación",2)
table(["ID","Supuesto","Impacto si es incorrecto","Acción"],[["AS-01","field1 se expresa en metros","Etiquetas y decisiones erróneas","Confirmar unidad"],["AS-02","field2 es lluvia en mm o evento binario","Representación incorrecta","Confirmar dominio"],["AS-03","field5 mide velocidad/variación del nivel","Interpretación equivocada","Definir fórmula y unidad"],["AS-04","field6 es comparable conceptualmente con el nivel","Lectura predictiva engañosa","Documentar modelo"],["AS-05","field7: 0 Normal, 1 Preventivo, 2 Alerta, 3 Crítico","Semáforo incorrecto","Validación prioritaria"],["AS-06","La zona horaria operativa es America/Lima","Fecha/hora desplazada","Confirmar zona"]],[1000,3150,3100,2110],8.2)
heading("6.3 Restricciones",2)
for text in ["Dependencia de disponibilidad, CORS y límites de la API de ThingSpeak.","Aplicación web sin backend en la primera versión.","No almacenar de forma permanente la Read API Key.","No alterar ni borrar los datos del canal.","El dashboard debe ser funcional con 50 a 500 lecturas por consulta.","La visualización requiere conexión a Internet para datos y biblioteca de gráficos."]: bullet(text)

functional = [
("RF-01","Configurar canal","El sistema debe permitir introducir y cambiar el ID del canal ThingSpeak desde la interfaz.","Alta","Propuesto"),("RF-02","Configurar clave de lectura","El sistema debe aceptar una Read API Key opcional y usarla únicamente para la sesión activa.","Alta","Propuesto"),("RF-03","Consultar feeds","El sistema debe solicitar los feeds JSON del canal seleccionado mediante HTTPS.","Alta","Implementado"),("RF-04","Normalizar datos","El sistema debe convertir fechas, identificadores y campos numéricos, preservando los valores ausentes como nulos.","Alta","Implementado"),("RF-05","Actualizar automáticamente","El sistema debe cargar la ventana completa una vez y consultar solo la lectura más reciente cada segundo mientras la página esté activa.","Alta","Implementado"),("RF-06","Actualizar manualmente","El usuario debe poder solicitar una actualización inmediata y recibir confirmación.","Media","Implementado"),("RF-07","Seleccionar ventana","El usuario debe elegir 50, 100, 250 o 500 lecturas.","Alta","Implementado"),("RF-08","Mostrar conectividad","El sistema debe indicar conexión pública/privada o falta de conexión.","Alta","Implementado"),("RF-09","Mostrar nivel actual","El sistema debe destacar el último nivel con unidad y diferencia frente a la lectura anterior.","Alta","Implementado"),("RF-10","Mostrar KPI","El sistema debe mostrar lluvia, temperatura, humedad y velocidad de la lectura más reciente.","Alta","Implementado"),("RF-11","Mostrar estado","El sistema debe traducir field7 a una etiqueta, descripción y color configurados.","Alta","Implementado con supuesto"),("RF-12","Graficar nivel observado","El sistema debe representar field1 en una serie temporal con escala propia.","Alta","Implementado"),("RF-13","Graficar predicción","El sistema debe representar field6 en una serie independiente para evitar aplanar field1.","Alta","Implementado"),("RF-14","Graficar velocidad","El sistema debe representar field5 en una escala independiente.","Media","Implementado"),("RF-15","Graficar clima","El sistema debe representar temperatura y humedad con dos ejes identificados.","Media","Implementado"),("RF-16","Graficar eventos","El sistema debe representar los códigos de estado a lo largo del tiempo.","Media","Implementado"),("RF-17","Generar resumen","El sistema debe calcular mínimo, máximo y promedio del nivel visible.","Media","Implementado"),("RF-18","Listar lecturas","El sistema debe mostrar las veinte lecturas más recientes en orden descendente.","Alta","Implementado"),("RF-19","Exportar CSV","El usuario debe poder descargar la ventana visible con fecha, entry_id y los siete campos.","Media","Implementado"),("RF-20","Manejar errores","El sistema debe conservar la interfaz y mostrar un mensaje cuando la API falle o no haya datos.","Alta","Implementado"),("RF-21","Adaptar la interfaz","El sistema debe reorganizar navegación, KPI, gráficos y tabla según el ancho disponible.","Alta","Implementado"),("RF-22","Formatear fecha local","El sistema debe presentar fechas en formato es-PE y zona local del navegador.","Media","Implementado"),("RF-23","No ocultar atípicos","Las escalas deben incluir los valores recibidos sin truncarlos silenciosamente.","Alta","Implementado"),("RF-24","Navegar por secciones","El usuario debe acceder a Resumen, Tendencias, Ambiente y Registros mediante navegación interna.","Baja","Implementado")]
functional.extend([
("RF-25","Seleccionar río","El usuario debe seleccionar un río del catálogo nacional sin confundir la identidad ni la telemetría de otra estación.","Alta","Implementado"),
("RF-26","Buscar río","El sistema debe filtrar el directorio nacional por nombre y mostrar hasta doce coincidencias relevantes.","Media","Implementado"),
("RF-27","Mostrar localidad","El sistema debe presentar río, localidad y región hidrográfica cuando la información haya sido validada.","Alta","Implementado parcial"),
("RF-28","Distinguir cobertura","El sistema debe diferenciar claramente una estación EN VIVO de un río con estación PENDIENTE.","Alta","Implementado"),
("RF-29","Configurar estación por río","El usuario debe asociar un canal ThingSpeak independiente al río seleccionado durante la sesión.","Alta","Implementado"),
("RF-30","Exportar contexto territorial","El CSV debe incluir río, localidad y región hidrográfica además de la telemetría.","Media","Implementado"),
("RF-31","Mostrar directorio nacional","El sistema debe informar el total de ríos catalogados y la cantidad de estaciones activas.","Media","Implementado"),
("RF-32","Evitar datos sustitutos","Un río sin estación no debe reutilizar ni mostrar las lecturas del Río Huallaga.","Alta","Implementado")])
heading("7. REQUISITOS FUNCIONALES",1)
add_p("Los requisitos se expresan de forma atómica y verificable. “Implementado” indica que existe una primera materialización en el prototipo; no equivale a aceptación del stakeholder.")
table(["ID","Nombre","Especificación","Prioridad","Estado"], functional,[850,1450,4700,1050,1310],7.7)

nonfunctional = [
("RNF-01","Rendimiento","La primera vista deberá mostrar datos en ≤ 3 s en una conexión estable de 10 Mbps, excluyendo indisponibilidad de ThingSpeak.","Prueba temporizada"),("RNF-02","Actualización","Una lectura disponible en la API deberá reflejarse en ≤ 25 s mientras la actualización automática esté activa.","Prueba de latencia"),("RNF-03","Seguridad","La Read API Key no deberá estar escrita en archivos fuente, URL fija, repositorio ni almacenamiento persistente.","Inspección técnica"),("RNF-04","Privacidad","La clave ingresada se almacenará solo en sessionStorage y se eliminará al cerrar la pestaña.","Prueba de sesión"),("RNF-05","Compatibilidad","El dashboard deberá funcionar en las dos últimas versiones estables de Edge, Chrome y Firefox.","Matriz de navegadores"),("RNF-06","Responsive","No deberá existir desplazamiento horizontal de la página entre 360 y 1920 px; la tabla puede usar desplazamiento interno.","Prueba de viewport"),("RNF-07","Accesibilidad","Navegación y controles operables por teclado, etiquetas asociadas y contraste AA en texto esencial.","Auditoría WCAG 2.2"),("RNF-08","Usabilidad","Un usuario nuevo deberá identificar nivel y estado actual en ≤ 10 s, sin asistencia.","Prueba de tarea"),("RNF-09","Fiabilidad","Los valores 0 y negativos válidos no se tratarán como ausentes; null y cadena vacía sí.","Pruebas unitarias"),("RNF-10","Mantenibilidad","La configuración de estado, canal y ventana deberá estar centralizada y documentada.","Revisión de código"),("RNF-11","Escalabilidad","La interfaz deberá representar 500 lecturas sin bloquear la interacción más de 200 ms después de recibir la respuesta.","Perfil de rendimiento"),("RNF-12","Trazabilidad","Cada requisito deberá relacionarse con al menos una historia o prueba.","Matriz de trazabilidad"),("RNF-13","Integridad visual","Ningún gráfico deberá mezclar magnitudes incompatibles sin ejes diferenciados o etiqueta explícita.","Revisión visual"),("RNF-14","Internacionalización","Textos y fechas visibles usarán español y formato regional es-PE.","Inspección UI"),("RNF-15","Recuperación","Tras un error temporal, la actualización manual o automática deberá restablecer el estado conectado sin recargar la página.","Prueba de recuperación")]
nonfunctional.extend([
("RNF-16","Integridad territorial","Ningún río podrá heredar el canal o la localidad de otro río salvo configuración explícita del usuario.","Prueba de selección"),
("RNF-17","Escalabilidad territorial","El selector deberá manejar al menos 900 nombres de ríos sin bloquear la interacción perceptiblemente.","Prueba con catálogo completo"),
("RNF-18","Despliegue continuo","Cada cambio aprobado en main deberá ejecutar pruebas y publicar GitHub Pages automáticamente.","Evidencia de GitHub Actions"),
("RNF-19","Disponibilidad pública","La URL principal deberá responder mediante HTTPS y no requerir autenticación.","Prueba HTTP externa"),
("RNF-20","Transparencia preventiva","La interfaz deberá declarar estación pendiente o datos no confirmados antes de presentar una conclusión operativa.","Revisión funcional")])
heading("8. REQUISITOS NO FUNCIONALES",1)
table(["ID","Atributo","Requisito medible","Verificación"], nonfunctional,[900,1400,5300,1760],7.8)

stories = [
("HU-01","Como operador, quiero ver el nivel y estado actual al abrir el dashboard para evaluar la estación rápidamente.",["Se muestran valor, unidad, estado y hora de actualización.","La lectura corresponde al último feed recibido.","Si no hay datos se informa el motivo."]),
("HU-02","Como analista, quiero cambiar la ventana de lecturas para revisar tendencias cortas o largas.",["Existen opciones 50/100/250/500.","Todos los gráficos, estadísticas y CSV usan la misma ventana."]),
("HU-03","Como analista, quiero ver nivel y predicción en escalas separadas para que los picos predictivos no oculten el comportamiento real.",["Cada variable dispone de su gráfico.","Se muestran todos los valores, incluidos atípicos."]),
("HU-04","Como responsable IoT, quiero configurar canal y clave sin editar el código para reutilizar la interfaz de manera segura.",["El canal es editable.","La clave es opcional y enmascarada.","No persiste al cerrar la pestaña."]),
("HU-05","Como operador, quiero reconocer estados por etiqueta y color para priorizar mi atención.",["Los códigos conocidos se traducen.","Un código desconocido se muestra como tal.","La información no depende solo del color."]),
("HU-06","Como analista, quiero consultar temperatura y humedad juntas para observar su relación.",["Se usan dos ejes identificados.","El tooltip muestra ambas magnitudes y unidades."]),
("HU-07","Como responsable del reporte, quiero exportar las lecturas a CSV para realizar análisis externo.",["Incluye fecha ISO, entry_id y siete campos.","El archivo corresponde a la ventana visible."]),
("HU-08","Como usuario móvil, quiero consultar KPI y tendencias sin perder información esencial.",["Funciona desde 360 px.","Los gráficos se adaptan.","La tabla usa scroll interno."]),
("HU-09","Como usuario, quiero saber si los datos están conectados para no interpretar información desactualizada como actual.",["Se indica conectado público/privado o sin conexión.","Se conserva la última hora recibida."]),
("HU-10","Como operador, quiero actualizar manualmente para comprobar un cambio sin esperar el intervalo automático.",["El botón inicia una consulta.","Muestra actividad y resultado."]),
("HU-11","Como Product Owner, quiero un resumen de mínimo, máximo y promedio para comprender la ventana sin calcularla manualmente.",["Cálculos ignoran null.","Los valores se actualizan al cambiar ventana."]),
("HU-12","Como usuario con teclado, quiero acceder a controles y secciones sin mouse para operar la interfaz de forma inclusiva.",["Orden de foco lógico.","Controles nativos activables por teclado.","Foco visible."])]
stories.extend([
("HU-13","Como operador nacional, quiero seleccionar el río que estoy supervisando para consultar únicamente su estación.",["El selector contiene el catálogo nacional.","El encabezado y la ficha territorial cambian con la selección.","La telemetría pertenece al canal asociado al río."]),
("HU-14","Como responsable IoT, quiero identificar los ríos que aún no tienen estación para planificar nuevas incorporaciones.",["El directorio diferencia EN VIVO y PENDIENTE.","Un río pendiente muestra valores vacíos y una instrucción de configuración."]),
("HU-15","Como analista, quiero buscar un río por nombre para encontrarlo sin recorrer cientos de opciones.",["La búsqueda ignora mayúsculas y minúsculas.","Los resultados se actualizan mientras se escribe.","Seleccionar un resultado actualiza todo el contexto."]),
("HU-16","Como responsable del reporte, quiero que el CSV conserve el contexto territorial para evitar mezclar registros de distintos ríos.",["Cada fila incluye río, localidad y región.","El nombre del archivo identifica el río y la fecha."])])
heading("9. HISTORIAS DE USUARIO Y CRITERIOS DE ACEPTACIÓN",1)
for sid, story, criteria in stories:
    heading(f"{sid}",2); add_p(story)
    p=doc.add_paragraph(); r=p.add_run("Criterios de aceptación"); set_run(r,size=9,color=TEAL,bold=True); p.paragraph_format.space_after=Pt(3)
    for c in criteria: bullet(c)

heading("10. CASOS DE USO",1)
cases=[("CU-01","Consultar dashboard","Operador","Canal configurado y conectividad","1) Abrir dashboard. 2) Consultar API. 3) Normalizar. 4) Mostrar KPI, estado y tendencias.","Información actual visible","Error controlado o ausencia de feeds"),("CU-02","Cambiar ventana","Analista","Dashboard cargado","1) Elegir cantidad. 2) Consultar nuevamente. 3) Recalcular todas las vistas.","Ventana consistente","API no disponible"),("CU-03","Configurar canal privado","Responsable IoT","Read API Key válida","1) Abrir configuración. 2) Ingresar canal y clave. 3) Conectar.","Conexión privada activa","401/404; clave no se persiste"),("CU-04","Exportar lecturas","Analista","Al menos una lectura","1) Seleccionar ventana. 2) Pulsar Exportar CSV. 3) Guardar archivo.","CSV descargado","Navegador bloquea descarga"),("CU-05","Recuperar conexión","Usuario","Existe un error temporal","1) Restablecer red/servicio. 2) Actualizar manualmente o esperar intervalo.","Estado conectado sin recargar","Error persiste y se informa")]
cases.extend([("CU-06","Seleccionar río","Operador nacional","Catálogo cargado","1) Abrir selector. 2) Elegir río. 3) Actualizar localidad, región y estación. 4) Consultar su canal si existe.","Contexto territorial consistente","Sin canal: estación pendiente"),("CU-07","Buscar río","Usuario","Directorio visible","1) Escribir nombre. 2) Revisar coincidencias. 3) Seleccionar resultado.","Río seleccionado","Sin coincidencias: mensaje informativo"),("CU-08","Vincular estación","Responsable IoT","Río seleccionado y canal conocido","1) Abrir configuración. 2) Ingresar ID y clave opcional. 3) Conectar. 4) Validar feeds.","Estación activa durante la sesión","Canal inválido o sin feeds")])
table(["ID","Caso","Actor","Precondición","Flujo principal","Resultado","Alternativa"],cases,[650,1200,1050,1450,2450,1300,1260],7.3)

heading("11. PRODUCT BACKLOG Y PRIORIZACIÓN",1)
backlog=[("PB-01","Conexión y normalización ThingSpeak","Debe","5","RF-01–RF-04"),("PB-02","Resumen y KPI en vivo","Debe","5","RF-08–RF-11"),("PB-03","Series con escalas correctas","Debe","8","RF-12–RF-16, RF-23"),("PB-04","Actualización automática/manual","Debe","3","RF-05–RF-06"),("PB-05","Ventanas de análisis","Debe","3","RF-07"),("PB-06","Tabla y trazabilidad","Debe","5","RF-18"),("PB-07","Exportación CSV","Debería","3","RF-19"),("PB-08","Responsive y navegación","Debe","5","RF-21, RF-24"),("PB-09","Mensajes y recuperación","Debe","5","RF-20"),("PB-10","Validación de estados/unidades","Debe","3","AS-01–AS-06"),("PB-11","Accesibilidad AA","Debería","5","RNF-07"),("PB-12","Alertas externas","Podría","13","Fuera de v1")]
backlog.extend([("PB-13","Selector y buscador nacional","Debe","8","RF-25–RF-26, RF-31"),("PB-14","Contexto territorial por río","Debe","5","RF-27–RF-28"),("PB-15","Canal independiente por río","Debe","8","RF-29, RF-32"),("PB-16","CSV con identidad territorial","Debería","3","RF-30"),("PB-17","Pipeline GitHub Pages","Debe","3","RNF-18–RNF-19")])
table(["ID","Elemento","MoSCoW","Puntos","Trazabilidad"],backlog,[900,3850,1300,1000,2310],8.2)
heading("11.1 Plan incremental sugerido",2)
table(["Sprint","Objetivo","Entregables","Salida de validación"],[["Sprint 1","Asegurar datos y semántica","Conexión, diccionario, prototipo de KPI","Unidades y estados aprobados"],["Sprint 2","Visualización a escala","Gráficos, filtros, resumen","Prueba con datos normales/atípicos"],["Sprint 3","Operación y calidad","Errores, CSV, responsive, a11y","Criterios RNF aprobados"],["Sprint 4","Aceptación","Correcciones y documentación","Acta de aceptación"]],[1150,2600,2900,2710],8.3)

heading("12. MATRIZ DE TRAZABILIDAD",1)
trace=[("NEC-01","Lectura inmediata del estado","RF-08–RF-11","HU-01, HU-05, HU-09","CP-01, CP-02"),("NEC-02","Escalas comprensibles","RF-12–RF-16, RF-23","HU-03, HU-06","CP-03, CP-04"),("NEC-03","Análisis por ventana","RF-07, RF-17","HU-02, HU-11","CP-05"),("NEC-04","Acceso seguro a canal privado","RF-01–RF-03","HU-04","CP-06, CP-07"),("NEC-05","Trazabilidad/exportación","RF-18, RF-19","HU-07","CP-08"),("NEC-06","Uso multidispositivo","RF-21, RF-24","HU-08, HU-12","CP-09, CP-10"),("NEC-07","Continuidad ante fallos","RF-20","HU-09, HU-10","CP-11"),("NEC-08","Actualidad de datos","RF-05, RF-06, RF-22","HU-01, HU-10","CP-12")]
trace.extend([("NEC-09","Cobertura nacional sin mezclar datos","RF-25, RF-28, RF-32","HU-13, HU-14","CP-13, CP-14"),("NEC-10","Localización y búsqueda","RF-26–RF-27, RF-31","HU-15","CP-15"),("NEC-11","Estaciones extensibles","RF-29","HU-13, HU-14","CP-16"),("NEC-12","Exportación territorial","RF-30","HU-16","CP-17"),("NEC-13","Publicación pública continua","RNF-18–RNF-19","Product Owner","CP-18")])
table(["Necesidad","Descripción","Requisitos","Historias","Pruebas"],trace,[1100,2700,2200,1900,1460],8.1)

heading("13. PLAN DE VALIDACIÓN Y ACEPTACIÓN",1)
tests=[("CP-01","Cargar canal 3420787","Se muestran entry_id y hora de la lectura más reciente"),("CP-02","Inyectar códigos 0–3 en datos controlados","Etiqueta textual y color coinciden con catálogo aprobado"),("CP-03","Usar nivel bajo y predicción atípica alta","Nivel conserva detalle en gráfico independiente"),("CP-04","Usar null, 0 y negativos","0/negativos válidos se muestran; null genera ausencia"),("CP-05","Cambiar 50 a 250 lecturas","Gráficos, resumen, tabla y CSV usan 250"),("CP-06","Buscar la clave en archivos y repositorio","No aparece ninguna clave real"),("CP-07","Cerrar y reabrir pestaña tras ingresar clave","sessionStorage ya no contiene la clave"),("CP-08","Exportar CSV","Columnas y cantidad coinciden con la ventana"),("CP-09","Viewport 360, 768, 1440 px","Sin scroll horizontal de página ni superposiciones"),("CP-10","Navegar solo con teclado","Todos los controles esenciales son alcanzables y accionables"),("CP-11","Simular 500 y luego 200 en API","Se informa error y luego se recupera sin recarga"),("CP-12","Publicar una lectura nueva","Aparece en ≤25 s con autoactualización")]
tests.extend([("CP-13","Seleccionar Río Huallaga","Se muestra canal 3420787, localidad y telemetría activa"),("CP-14","Seleccionar un río sin canal","No se reutilizan lecturas del Huallaga y aparece Estación pendiente"),("CP-15","Buscar Marañón o Rímac","El río aparece entre las coincidencias y puede seleccionarse"),("CP-16","Asignar un canal de prueba a otro río","La configuración afecta solo a ese río y solo durante la sesión"),("CP-17","Exportar CSV del Huallaga","El archivo incluye río, localidad y región en cada fila"),("CP-18","Enviar un cambio aprobado a main","GitHub Actions ejecuta pruebas y GitHub Pages responde HTTP 200")])
table(["ID","Escenario","Resultado esperado"],tests,[1000,3600,4760],8.3)
heading("13.1 Criterio de salida",2)
for text in ["Todos los requisitos de prioridad Alta aceptados o con excepción aprobada.","Cero defectos críticos o bloqueantes abiertos.","Unidades y estados confirmados por el responsable IoT.","CP-01 a CP-18 ejecutados; al menos 95 % aprobados y 100 % de los críticos aprobados.","Product Owner y usuario operativo firman el acta de aceptación."]: bullet(text)

heading("14. RIESGOS",1)
risks=[("R-01","Interpretación incorrecta de field7","Alta","Alta","Crítico","Validar catálogo antes de operar"),("R-02","Unidades no confirmadas","Alta","Alta","Crítico","Aprobar diccionario de datos"),("R-03","Exposición de API Key","Media","Alta","Alto","Solo sesión; revisión de código"),("R-04","Caída o límite de ThingSpeak","Media","Media","Medio","Mensajes y recuperación"),("R-05","Atípicos aplastan gráficos","Alta","Media","Alto","Escalas independientes"),("R-06","Datos desactualizados se interpretan como actuales","Media","Alta","Alto","Hora visible y estado de conexión"),("R-07","CORS/CDN indisponible","Baja","Media","Bajo","Documentar dependencia; opción de empaquetado futuro"),("R-08","Color no accesible","Media","Media","Medio","Etiqueta textual + contraste")]
risks.extend([("R-09","Asignar una localidad no validada","Media","Alta","Alto","Mostrar Por validar y aprobar catálogo territorial"),("R-10","Mezclar canales entre ríos","Baja","Alta","Alto","Configuración por nombre y prueba CP-14"),("R-11","Catálogo IGN cambia o deja de responder","Baja","Media","Bajo","Mantener copia versionada y sincronización controlada"),("R-12","Falla del despliegue continuo","Baja","Media","Bajo","Conservar versión anterior y revisar GitHub Actions")])
table(["ID","Riesgo","Prob.","Impacto","Nivel","Mitigación"],risks,[750,3150,900,950,1000,2610],8.2)

heading("15. GLOSARIO Y ANEXOS",1)
heading("15.1 Glosario",2)
table(["Término","Definición"],[["Dashboard","Interfaz que sintetiza indicadores, tendencias y registros."],["Feed","Registro individual almacenado por ThingSpeak."],["Read API Key","Credencial destinada a consultar un canal privado."],["KPI","Indicador clave mostrado para lectura rápida."],["ThingSpeak","Servicio IoT de MathWorks usado como fuente de datos."],["Telemetría","Datos medidos y enviados remotamente por la estación."],["Valor atípico","Observación alejada del patrón general que debe conservarse y contextualizarse."],["Responsive","Capacidad de adaptar la composición a diferentes tamaños de pantalla."],["sessionStorage","Almacenamiento del navegador limitado a la pestaña/sesión."]],[2200,7160],8.7)
heading("15.2 Evidencia de lectura observada",2)
table(["Fecha UTC","entry_id","Nivel","Lluvia","Temp","Hum","Velocidad","Predicción","Estado"],[["16/07/2026 14:11:56","269","2.45898","0","23.6","78.8","0","2.45898","0"],["16/07/2026 14:12:17","270","4.78795","0","23.6","78.8","59.99073","184.76013","3"],["16/07/2026 14:12:36","271","2.46241","0","23.7","78.6","0","2.46241","0"],["16/07/2026 14:12:56","272","0.50045","0","23.7","78.6","0","0.50045","0"],["16/07/2026 14:13:16","273","0.51074","0","23.8","78.5","0.10288","0.81939","0"]],[1750,800,850,700,700,700,1100,1100,660],7.4)
add_p("Fuente: consulta de solo lectura al endpoint público del canal 3420787 durante la elicitación. La muestra se incluye para trazabilidad; el dashboard consulta los datos vigentes en cada sesión.")
heading("15.3 Preguntas abiertas",2)
for text in ["¿Cuáles son las unidades oficiales de Nivel, Lluvia, Velocidad y Nivel Predicción?","¿Qué fórmula produce field5 y field6?","¿Cuál es el catálogo definitivo y la acción asociada a cada valor de field7?","¿Qué umbrales convierten un valor en preventivo, alerta o crítico?","¿El canal continuará público en producción?","¿Cuál es la zona horaria que debe mostrarse?","¿Se requieren alertas externas o historial de auditoría en una segunda versión?","¿Cuál es la política de retención y respaldo de los datos?"]: bullet(text)
callout("Próximo paso", "Realizar una sesión de validación de 45 minutos con Product Owner, responsable IoT y usuario operativo; resolver AS-01 a AS-06 y actualizar el estado de los requisitos de Propuesto a Aprobado.")

heading("16. DESPLIEGUE, DOMINIO Y OPERACIÓN",1)
add_p("El Dashboard Hidrológico de Prevención se encuentra publicado con acceso público y HTTPS. La URL principal es https://jaqzo8.github.io/Dashboard_Hidrol-gicoIoT/ y se genera desde el repositorio https://github.com/JaqzO8/Dashboard_Hidrol-gicoIoT mediante GitHub Actions.")
heading("16.1 Costo y sostenibilidad",2)
table(["Concepto","Costo","Observación"],[["Repositorio público GitHub","S/ 0","Código y documentación versionados"],["GitHub Pages","S/ 0","Hosting estático para el alcance actual"],["Certificado HTTPS","S/ 0","Administrado automáticamente"],["Servidor propio","S/ 0","No requerido"],["Despliegue alternativo","S/ 0","OpenAI Sites como respaldo"]],[3200,1400,4760],8.5)
add_p("La alternativa es rentable para demostración académica y operación de bajo tráfico porque evita compra de dominio, servidor y certificado. Un dominio institucional propio podrá incorporarse posteriormente si se necesita identidad de marca.")
heading("16.2 Arquitectura y seguridad",2)
for text in ["GitHub Pages entrega una aplicación compilada y optimizada con Vite; el código fuente está separado en interfaz, sincronización en vivo y cliente ThingSpeak.","El cliente realiza una carga completa inicial y luego consulta el endpoint de última entrada cada segundo para actualizar solo lo necesario; la llegada de valores nuevos depende de la frecuencia de publicación del sensor.","El canal público no requiere credenciales incorporadas en el código ni en el hosting.","Una Read API Key opcional se guarda únicamente en sessionStorage y desaparece al cerrar la pestaña.","El dashboard es de solo lectura: no modifica ni elimina registros IoT."]: bullet(text)
heading("16.3 Evidencia del despliegue",2)
table(["Elemento","Resultado"],[["URL principal","https://jaqzo8.github.io/Dashboard_Hidrol-gicoIoT/"],["Repositorio","https://github.com/JaqzO8/Dashboard_Hidrol-gicoIoT"],["Acceso","Público, sin inicio de sesión"],["Protocolo","HTTPS"],["Automatización","GitHub Actions: pruebas y despliegue desde main"],["Fuente IoT inicial","ThingSpeak, canal 3420787 - Río Huallaga"],["Actualización","Carga inicial y consulta incremental cada segundo"]],[2600,6760],8.3)
heading("16.4 Operación durante la primera semana",2)
for text in ["Confirmar que la URL principal responde mediante HTTPS.","Comparar la lectura más reciente con el canal 3420787 y verificar que la marca de tiempo avance.","Probar selección del Huallaga y de un río con estación pendiente.","Probar actualización manual, selector de registros, vista móvil y exportación CSV.","Revisar el último flujo Deploy GitHub Pages y conservar main en estado desplegable.","Si una publicación introduce una falla, volver a desplegar la última versión aprobada."]: bullet(text)
callout("Criterio de permanencia", "No existe tarea de caducidad configurada. La URL continuará activa mientras el repositorio y GitHub Pages permanezcan habilitados y no se incumplan las políticas del proveedor.")

heading("17. DISEÑO DETALLADO DE LA INTERFAZ",1)
heading("17.1 Sistema visual",2)
table(["Elemento","Especificación","Propósito"],[["Paleta base","Azul petróleo y fondos oscuros","Reducir fatiga y comunicar contexto técnico"],["Acento principal","Turquesa","Destacar telemetría activa y acciones"],["Estados","Verde, amarillo, naranja y rojo con texto","Comunicar severidad sin depender solo del color"],["Tipografía","Manrope para cuerpo y Space Grotesk para títulos","Jerarquía contemporánea y alta legibilidad"],["Contenedores","Paneles con bordes suaves y profundidad moderada","Separar información sin fragmentarla"],["Movimiento","Transiciones breves y actualización visible","Confirmar interacción sin distraer"]],[1800,3700,3860],8.4)
heading("17.2 Componentes de la pantalla",2)
table(["Componente","Contenido","Interacción","Estado alternativo"],[["Barra lateral","Marca, navegación, río y canal activo","Navegar a secciones y abrir configuración","Oculta parcialmente en móvil"],["Encabezado","Título, selector de río, ventana y actualización","Cambiar contexto, rango o refrescar","Hora vacía sin datos"],["Panel preventivo","Nivel actual, tendencia y estado","Lectura inmediata","Sin estación o sin conexión"],["KPI","Nivel, lluvia, temperatura, humedad y velocidad","Lectura comparativa","Guion cuando el dato es nulo"],["Gráfico de nivel","Serie field1 con escala propia","Tooltip temporal","Panel vacío sin feeds"],["Predicción","Serie field6 separada","Comparar tendencia prevista","Unidad marcada por validar"],["Velocidad","Serie field5 independiente","Observar cambios bruscos","Unidad marcada por validar"],["Ambiente","Temperatura y humedad con dos ejes","Comparación por fecha","Conexión de valores nulos"],["Eventos","Códigos field7 por lectura","Reconocer frecuencia de estados","Código desconocido visible"],["Red nacional","Conteos, buscador y directorio de ríos","Buscar y seleccionar","Estación pendiente"],["Registros","Veinte lecturas recientes","Desplazar y exportar CSV","Mensaje sin datos"],["Diálogo de conexión","Río, localidad, canal y clave opcional","Asociar estación durante la sesión","Error de API informado"]],[1650,2850,2350,2510],7.7)
heading("17.3 Comportamiento responsive",2)
table(["Rango","Composición esperada","Controles críticos"],[["≥ 1100 px","Sidebar fija, hero de tres áreas, KPI en cinco columnas","Selector, rango, refrescar, CSV"],["761-1099 px","Sidebar compacta, hero reorganizado, KPI en tres columnas","Todos visibles"],["≤ 760 px","Cabecera móvil, paneles en una columna, tabla con scroll interno","Selector, rango y configuración"],["≤ 440 px","KPI en una columna y medidor reducido","Sin pérdida de información esencial"]],[1500,5000,2860],8.4)

heading("18. MODELO MULTI-RÍO Y PREVENCIÓN",1)
heading("18.1 Catálogo territorial",2)
add_p("El catálogo contiene 969 nombres de ríos obtenidos de la capa Ríos y Quebradas de la cartografía oficial IDEP/IGN a escala 1:500 000. Su función es ofrecer cobertura de selección; la existencia en el catálogo no implica que haya una estación IoT activa ni que la localidad esté validada para uso operativo.")
table(["Estado","Condición","Presentación","Acción permitida"],[["EN VIVO","Existe canal configurado","KPI, gráficos y registros","Consultar y exportar"],["PENDIENTE","No existe canal configurado","Valores vacíos y mensaje explícito","Configurar estación"],["POR VALIDAR","Falta confirmar localidad, región o unidades","Etiqueta de advertencia","Validación documental"],["SIN CONEXIÓN","La estación existe pero la API falla","Última hora y error controlado","Reintentar manual/automático"]],[1500,2400,3000,2460],8.3)
heading("18.2 Estación inicial: Río Huallaga",2)
table(["Atributo","Definición"],[["Río","Río Huallaga"],["Canal ThingSpeak","3420787"],["Región hidrográfica","Amazonas"],["Ámbito territorial mostrado","Huánuco, San Martín y Loreto"],["Estado de integración","EN VIVO"],["Frecuencia de consulta incremental","1 segundo"],["Modo","Solo lectura"]],[2800,6560],8.6)
heading("18.3 Reglas preventivas",2)
for text in ["Toda alerta debe conservar fecha y hora de la lectura que la originó.","El estado textual debe acompañar siempre al color.","La ausencia de estación o de datos nunca se interpretará como condición Normal.","Los valores atípicos se conservarán y se mostrarán en una escala que permita identificarlos.","Las conclusiones preventivas dependen de unidades, umbrales y códigos formalmente aprobados.","El dashboard apoya la observación; la activación de protocolos oficiales corresponde a la autoridad responsable."]: bullet(text)

heading("19. IMPLEMENTACIÓN, DEVOPS Y MANTENIMIENTO",1)
heading("19.1 Estructura técnica",2)
table(["Archivo / área","Responsabilidad"],[["index.html","Punto de montaje y estructura semántica"],["styles.css","Sistema visual y comportamiento responsive"],["app.js","Coordinación de interfaz, estado, gráficos y exportación"],["src/services/thingspeak.js","Cliente HTTPS, validación y consultas completa/incremental"],["src/core/live-feed.js","Consulta cada segundo, cancelación, pausa y reintentos"],["vite.config.js","Compilación modular y optimización de producción"],["data/peru-rivers.js","Catálogo nacional versionado"],["tests/dashboard.test.mjs","Pruebas unitarias de datos y configuración territorial"],[".github/workflows/pages.yml","Pruebas y publicación automática en GitHub Pages"],["docs/GUIA_DESPLIEGUE.md","Operación, seguridad, recuperación y costos"]],[3000,6360],8.5)
heading("19.2 Flujo de entrega continua",2)
for text in ["El equipo confirma el alcance del cambio y actualiza código, pruebas y documentación.","Se ejecutan las pruebas locales y se revisa que no existan claves en el repositorio.","El cambio aprobado se integra en la rama main.","GitHub Actions descarga el código, configura Node.js y ejecuta npm test.","Si las pruebas aprueban, Vite compila módulos, optimiza los recursos y GitHub Pages publica el resultado versionado.","Se verifica la URL pública, el Río Huallaga y la conexión de ThingSpeak.","Ante una regresión se revierte o vuelve a desplegar la última revisión aprobada."]: bullet(text)
heading("19.3 Matriz de mantenimiento",2)
table(["Frecuencia","Actividad","Evidencia","Responsable"],[["Diaria durante piloto","Revisar actualidad de la última lectura","Hora y entry_id","Operador"],["Semanal","Probar selector, actualización y CSV","Checklist firmado","Responsable IoT"],["Mensual","Revisar dependencias CDN y límites de API","Registro técnico","Desarrollo"],["Por nueva estación","Validar río, localidad, canal, unidades y estados","Ficha de alta","IoT + Product Owner"],["Por cambio de código","Ejecutar pruebas y confirmar GitHub Pages","Flujo exitoso","Desarrollo"],["Trimestral","Revisar requisitos, riesgos y preguntas abiertas","Nueva versión del documento","Comité"]],[1550,3300,2500,2010],8.2)
heading("19.4 Definición de terminado",2)
for text in ["Código y documentación actualizados en el repositorio público.","Pruebas automatizadas aprobadas y URL pública disponible por HTTPS.","Sin API Keys, contraseñas o tokens en archivos versionados.","Requisitos, historia y prueba relacionados para cada cambio funcional.","Validación responsive y de teclado sin defectos críticos.","Responsable IoT confirma unidades, estado y canal de cualquier estación nueva.","Product Owner acepta el resultado o registra una excepción explícita."]: bullet(text)
callout("Resultado esperado", "El Dashboard Hidrológico de Prevención debe comunicar qué río se observa, qué estación aporta los datos, cuándo se actualizó y qué condición preventiva se interpreta, sin ocultar incertidumbre ni mezclar fuentes territoriales.")

# Final formatting pass
for section in doc.sections: setup_header_footer(section)
for p in doc.paragraphs:
    if p.style.name == "Normal" and p.text.strip(): p.paragraph_format.widow_control = True

OUT.parent.mkdir(parents=True, exist_ok=True)
doc.save(OUT)
print(OUT)

